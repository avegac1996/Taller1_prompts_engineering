"""Genera el informe completo en Word (.docx) a partir de los archivos del proyecto."""

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

PROJECT = Path(__file__).parent
MD_OUTPUT = PROJECT / "informe_completo.md"
DOCX_OUTPUT = PROJECT / "Informe_Taller1_PromptsEngineering_generando.docx"
DOCX_FINAL = PROJECT / "Informe_Taller1_PromptsEngineering.docx"


def read_file(path: Path) -> str:
    if not path.exists():
        return f"*Archivo {path.name} no encontrado.*"
    return path.read_text(encoding="utf-8")


def add(lines, text=""):
    lines.append(text)


def add_heading(lines, text, level=1):
    add(lines, f"{'#' * level} {text}")
    add(lines, "")


def add_para(lines, text):
    add(lines, text)
    add(lines, "")


def add_code(lines, text, lang=""):
    # Usa 4 backticks para evitar conflictos con triple backticks internos.
    add(lines, f"````{lang}")
    add(lines, text)
    add(lines, "````")
    add(lines, "")


def add_centered_paragraph(doc, text, font_size=12, bold=False, font_name="Calibri"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name
    return p


def insert_paragraph_at_beginning(doc, paragraph):
    body = doc._element.body
    body.remove(paragraph._element)
    body.insert(0, paragraph._element)


def insert_paragraph_after(target, paragraph):
    body = target._element.getparent()
    body.remove(paragraph._element)
    for i, child in enumerate(body):
        if child is target._element:
            body.insert(i + 1, paragraph._element)
            return
    body.append(paragraph._element)


def add_cover(doc):
    """Inserta una carátula al inicio del documento."""
    # Se agregan en orden inverso para que queden en el orden correcto al insertar al inicio.
    # El último en insertarse (título) quedará primero; el primero en insertarse (fecha) quedará último.
    p_date = add_centered_paragraph(doc, "Julio de 2026", 12)
    insert_paragraph_at_beginning(doc, p_date)

    insert_paragraph_at_beginning(doc, doc.add_paragraph())

    p_author = add_centered_paragraph(doc, "Estudiante: Jhossua Vega", 12)
    insert_paragraph_at_beginning(doc, p_author)

    insert_paragraph_at_beginning(doc, doc.add_paragraph())

    p_project = add_centered_paragraph(doc, "Aplicación profesional de Lista de Tareas en Python", 14)
    insert_paragraph_at_beginning(doc, p_project)

    insert_paragraph_at_beginning(doc, doc.add_paragraph())

    p_subject = add_centered_paragraph(doc, "Taller 1: Prompts Engineering", 18)
    insert_paragraph_at_beginning(doc, p_subject)

    insert_paragraph_at_beginning(doc, doc.add_paragraph())

    p_title = add_centered_paragraph(doc, "Informe de entrega", 28, bold=True)
    insert_paragraph_at_beginning(doc, p_title)

    # El salto de página va al final del último elemento de la carátula (fecha).
    p_date.runs[0].add_break(WD_BREAK.PAGE)

    return p_date


def add_introduction(doc, last_cover_para):
    """Inserta la introducción después de la carátula."""
    # Se agregan en orden inverso para que queden en el orden correcto al insertar después del último párrafo de la carátula.
    p_break = doc.add_paragraph()
    p_break.add_run().add_break(WD_BREAK.PAGE)
    insert_paragraph_after(last_cover_para, p_break)

    p3 = doc.add_paragraph(
        "El resultado final es un programa modular, con una clase TaskManager que gestiona la persistencia JSON, y una interfaz CLI basada en argparse. El proyecto incluye pruebas unitarias, validaciones, manejo de errores, exportación CSV y está alojado en GitHub para su revisión."
    )
    insert_paragraph_after(last_cover_para, p3)

    p2 = doc.add_paragraph(
        "El trabajo se organizó en cuatro prompts estructurados que cubren: arquitectura y pruebas, interfaz de línea de comandos, funcionalidades avanzadas y filtros, y finalmente UX, colores y documentación. Cada prompt incluye secciones de rol, contexto, tarea, restricciones, criterios de aceptación QA y consideraciones UX."
    )
    insert_paragraph_after(last_cover_para, p2)

    p1 = doc.add_paragraph(
        "Este informe presenta el desarrollo de una aplicación profesional de Lista de Tareas en Python, ejecutable desde la terminal. El objetivo fue aplicar técnicas de prompt engineering para generar un entregable robusto, ordenado y con enfoque de calidad (QA) y experiencia de usuario (UX)."
    )
    insert_paragraph_after(last_cover_para, p1)

    p_heading = doc.add_heading("Introducción", level=1)
    insert_paragraph_after(last_cover_para, p_heading)


def style_terminal_blocks(doc):
    """Aplica fondo negro y texto blanco a todos los bloques de código (estilo Source Code)."""
    for paragraph in doc.paragraphs:
        if paragraph.style.name != "Source Code":
            continue
        pPr = paragraph._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "000000")
        pPr.append(shd)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)


def format_docx(path):
    """Post-procesa el .docx para agregar carátula, introducción y estilo de terminal."""
    doc = Document(path)
    last_cover_para = add_cover(doc)
    add_introduction(doc, last_cover_para)
    style_terminal_blocks(doc)
    doc.save(path)


def main():
    lines = []

    # La carátula y la introducción se agregan después con python-docx.
    add_heading(lines, "1. Datos generales", level=2)
    add_para(lines, "- **Asignatura:** Taller 1: Prompts Engineering")
    add_para(lines, "- **Tema:** Crear una aplicación de Lista de Tareas en Python desde la terminal utilizando prompt engineering.")
    add_para(lines, "- **Nombre del estudiante:** Jhossua Vega")
    add_para(lines, "- **Fecha de entrega:** julio de 2026")

    add_heading(lines, "2. Proveedor y modelo de IA", level=2)
    add_para(lines, "- **Proveedor:** Cascade (asistente de código integrado en el IDE)")
    add_para(lines, "- **Modelo de IA:** OpenAI GPT-4o")

    add_heading(lines, "3. Repositorio en GitHub", level=2)
    add_para(lines, "URL: https://github.com/avegac1996/Taller1_prompts_engineering.git")

    add_heading(lines, "4. Prompts utilizados", level=2)
    add_para(lines, "Se diseñaron cuatro prompts estructurados con enfoque de QA y UX. El detalle completo se presenta a continuación.")
    add(lines, read_file(PROJECT / "prompts.md"))
    add(lines, "")

    add_heading(lines, "5. Resultado final", level=2)
    add_para(lines, "La aplicación entregada es una Lista de Tareas profesional que se ejecuta completamente desde la terminal. Incluye capa de dominio (clase `TaskManager`), interfaz CLI con `argparse`, validaciones, filtros, estadísticas, exportación CSV y una suite de pruebas unitarias.")

    add_heading(lines, "5.1. Código fuente - todo.py", level=3)
    add_para(lines, "El siguiente listado muestra el archivo principal de la aplicación:")
    add_code(lines, read_file(PROJECT / "todo.py"), lang="python")

    add_heading(lines, "5.2. Pruebas unitarias - tests/test_todo.py", level=3)
    add_para(lines, "Se implementó una suite de 23 pruebas unitarias que cubren validaciones, CRUD, búsqueda, filtros, ordenamiento, detección de tareas vencidas, exportación CSV y persistencia.")
    add_code(lines, read_file(PROJECT / "tests" / "test_todo.py"), lang="python")

    add_heading(lines, "6. Evidencia de ejecución", level=2)
    add_para(lines, "A continuación se incluyen las capturas de las pruebas manuales realizadas en la terminal. Los bloques simulan el fondo negro y la tipografía de consola.")
    add(lines, read_file(PROJECT / "capturas.md"))
    add(lines, "")

    add_heading(lines, "7. Documentación del proyecto", level=2)

    add_heading(lines, "7.1. README.md", level=3)
    add(lines, read_file(PROJECT / "README.md"))
    add(lines, "")

    add_heading(lines, "7.2. requirements.txt", level=3)
    add_code(lines, read_file(PROJECT / "requirements.txt"), lang="text")

    add_heading(lines, "7.3. .gitignore", level=3)
    add_code(lines, read_file(PROJECT / ".gitignore"), lang="text")

    add_heading(lines, "8. Plan de calidad y conclusiones", level=2)
    add_para(lines, "1. **Arquitectura:** La lógica de negocio está separada de la interfaz CLI mediante la clase `TaskManager`, facilitando las pruebas y el mantenimiento.")
    add_para(lines, "2. **Persistencia:** Las tareas se almacenan en un archivo JSON en el directorio personal del usuario (`%USERPROFILE%\\.todo_cli\\tasks.json`). El sistema se recupera automáticamente ante archivos corruptos.")
    add_para(lines, "3. **Validación:** Se validan prioridades (`baja`, `media`, `alta`) y fechas en formato ISO (`AAAA-MM-DD`). Los errores se reportan por `stderr` con códigos de salida distintos de cero.")
    add_para(lines, "4. **Pruebas:** 23 pruebas unitarias se ejecutan exitosamente con `python -m unittest discover tests`.")
    add_para(lines, "5. **UX:** Colores semánticos opcionales, tabla alineada, fechas vencidas resaltadas y mensajes de error informativos.")
    add_para(lines, "6. **Entrega:** El proyecto se encuentra en GitHub y es ejecutable en Python 3.14 desde la terminal.")

    MD_OUTPUT.write_text("\n".join(lines), encoding="utf-8")

    subprocess.run(
        [
            "pandoc",
            str(MD_OUTPUT),
            "-o",
            str(DOCX_OUTPUT),
            "--from=markdown",
            "--to=docx",
        ],
        check=True,
    )

    print(f"Documento base generado: {DOCX_OUTPUT}")
    format_docx(DOCX_OUTPUT)

    try:
        if DOCX_FINAL.exists():
            DOCX_FINAL.unlink()
        DOCX_OUTPUT.rename(DOCX_FINAL)
        print(f"Documento final: {DOCX_FINAL}")
    except OSError as exc:
        print(f"No se pudo reemplazar el archivo anterior (puede estar abierto): {exc}")
        print(f"Documento final: {DOCX_OUTPUT}")


if __name__ == "__main__":
    main()
